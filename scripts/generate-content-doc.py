"""
Generates CONTENT_UPDATES.docx with bold titles and spacing between items.
Run: python3 scripts/generate-content-doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── styles ────────────────────────────────────────────────

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return p

def spacer():
    doc.add_paragraph('')
    doc.add_paragraph('')

def item_title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p

def field(label, value):
    p = doc.add_paragraph()
    run_label = p.add_run(f'{label}: ')
    run_label.bold = False
    run_label.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    p.add_run(value)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    return p

# ── intro ─────────────────────────────────────────────────

heading('MIND Lab — Website Content Updates')
doc.add_paragraph(
    'To add content to the website, fill in the sections below and share '
    'the updated doc with the webmaster. Just follow the same format as the existing examples.'
)

# ══════════════════════════════════════════════════════════
#  NEWS
# ══════════════════════════════════════════════════════════

doc.add_page_break()
heading('NEWS')
doc.add_paragraph(
    'Add new items at the top. Date format: YYYY-MM-DD\n'
    'Types: paper · award · grant · event · press · misc'
)
doc.add_paragraph('')

news = [
    {
        'title': 'Two Papers Accepted at ICWSM 2026',
        'date': '2026-06-01',
        'type': 'paper',
        'description': 'Two papers accepted at ICWSM 2026: "Can Large Language Models Assess the Social Impact of Conspiracy Theories?" and "Multimodal Large Language Models as Synthetic Participants in Video-Based Studies: An Evaluation" (SocialLLM Workshop).',
        'link': '',
    },
    {
        'title': 'Dr. Zhou Co-Chairs WSDM 2026 Doctoral Consortium',
        'date': '2026-02-22',
        'type': 'event',
        'description': 'Dr. Zhou served as Co-Chair of the Doctoral Consortium at WSDM 2026, held in Boise, Idaho.',
        'link': 'https://wsdm-conference.org/2026/index.php/call-for-wsdm-doctoral-consortium/',
    },
    {
        'title': 'Paper Published in Scientific Reports',
        'date': '2026-01-01',
        'type': 'paper',
        'description': "Our paper \"Users' Prompting Strategies and ChatGPT's Contextual Adaptation Shape Conversational Information-Seeking Experiences\" has been published in Scientific Reports and presented at ICA 2026.",
        'link': '',
    },
    {
        'title': 'New Preprint: ChatBot Influence in Value-Driven Decision Making',
        'date': '2025-11-19',
        'type': 'paper',
        'description': 'New preprint: "A Crowdsourced Study of ChatBot Influence in Value-Driven Decision Making Scenarios" now available on arXiv.',
        'link': 'https://arxiv.org/abs/2511.15857',
    },
    {
        'title': 'PaperPing Presented at CSCW Companion 2025',
        'date': '2025-10-01',
        'type': 'paper',
        'description': 'Our paper "PaperPing: A Socially-aware AI Agent that Recommends Academic Papers to Research Group Chats with Contextualized Explanations" was presented at CSCW Companion 2025.',
        'link': '',
    },
    {
        'title': "New Preprint: ChatGPT User Prompting Strategies",
        'date': '2025-09-29',
        'type': 'paper',
        'description': "New preprint: \"Users' Prompting Strategies and ChatGPT's Contextual Adaptation Shape Conversational Information-Seeking Experiences\" now available on arXiv.",
        'link': 'https://arxiv.org/abs/2509.25513',
    },
    {
        'title': 'Dr. Zhou Joins Boise State University as Assistant Professor',
        'date': '2025-09-29',
        'type': 'misc',
        'description': 'Dr. Xinyi Zhou has joined the Department of Computer Science at Boise State University as an Assistant Professor, where she directs the MIND Lab.',
        'link': 'https://www.boisestate.edu/coen-cs/2025/09/29/computer-science-welcomes-new-faculty-fall-2025/',
    },
]

for i, n in enumerate(news):
    item_title(n['title'])
    field('Date', n['date'])
    field('Type', n['type'])
    field('Description', n['description'])
    field('Link', n['link'] or '(leave blank if none)')
    if i < len(news) - 1:
        spacer()

# ══════════════════════════════════════════════════════════
#  PUBLICATIONS
# ══════════════════════════════════════════════════════════

doc.add_page_break()
heading('PUBLICATIONS')
doc.add_paragraph(
    'Add new items at the top. Types: conference · journal · workshop · preprint\n'
    'Authors: comma-separated, exactly as they should appear. Lab members are bolded automatically.\n'
    'Featured: yes or no (pick at most 3 total)'
)
doc.add_paragraph('')

pubs = [
    {
        'title': 'Multimodal Large Language Models as Synthetic Participants in Video-Based Studies: An Evaluation',
        'authors': 'Prabal Shrestha, Bohan Jiang, Haoning Xue, Huan Liu, Xinyi Zhou',
        'venue': 'Workshop on Large Language Models for Social Reasoning and Simulation (SocialLLM) @ ICWSM',
        'year': '2026',
        'type': 'workshop',
        'pdf': '',
        'link': '',
        'tags': 'Multimodal ML, LLMs, Human Simulation, Video Understanding',
        'featured': 'yes',
    },
    {
        'title': 'Can Large Language Models Assess the Social Impact of Conspiracy Theories?',
        'authors': 'Bohan Jiang, Dawei Li, Zhen Tan, Xinyi Zhou, Ashwin Rao, Kristina Lerman, H. Russell Bernard, Huan Liu',
        'venue': 'AAAI International Conference on Web and Social Media (ICWSM)',
        'year': '2026',
        'type': 'conference',
        'pdf': 'https://arxiv.org/abs/2412.07019',
        'link': '',
        'tags': 'LLMs, Misinformation, Social Media, Conspiracy Theories',
        'featured': 'no',
    },
    {
        'title': "Users' Prompting Strategies and ChatGPT's Contextual Adaptation Shape Conversational Information-Seeking Experiences",
        'authors': 'Haoning Xue, Yoo Jung Oh, Xinyi Zhou, Xinyu Zhang, Berit Oxley',
        'venue': 'Scientific Reports',
        'year': '2026',
        'type': 'journal',
        'pdf': 'https://arxiv.org/abs/2509.25513',
        'link': '',
        'tags': 'LLMs, Human-AI Interaction, Information Seeking',
        'featured': 'no',
    },
    {
        'title': 'A Crowdsourced Study of ChatBot Influence in Value-Driven Decision Making Scenarios',
        'authors': 'Anthony Wise, Xinyi Zhou, Martin Reimann, Anind Dey, Leilani Battle',
        'venue': 'arXiv preprint',
        'year': '2025',
        'type': 'preprint',
        'pdf': 'https://arxiv.org/abs/2511.15857',
        'link': '',
        'tags': 'LLMs, Human-AI Interaction, Decision Making',
        'featured': 'yes',
    },
    {
        'title': 'PaperPing: A Socially-aware AI Agent that Recommends Academic Papers to Research Group Chats with Contextualized Explanations',
        'authors': 'Ruotong Wang, Xinyi Zhou, Lin Qiu, Joseph Chee Chang, Jonathan Bragg, Amy X. Zhang',
        'venue': 'Companion Proceedings of the ACM Conference on Computer-Supported Cooperative Work and Social Computing (CSCW Companion)',
        'year': '2025',
        'type': 'workshop',
        'pdf': '',
        'link': 'https://dl.acm.org/doi/10.1145/3715070.3757230',
        'tags': 'LLM Agents, Social Computing, Human-AI Interaction',
        'featured': 'no',
    },
]

for i, p in enumerate(pubs):
    item_title(p['title'])
    field('Authors', p['authors'])
    field('Venue', p['venue'])
    field('Year', p['year'])
    field('Type', p['type'])
    field('PDF', p['pdf'] or '(leave blank if none)')
    field('Link', p['link'] or '(leave blank if none)')
    field('Tags', p['tags'])
    field('Featured', p['featured'])
    if i < len(pubs) - 1:
        spacer()

# ══════════════════════════════════════════════════════════
#  RESEARCH AREAS
# ══════════════════════════════════════════════════════════

doc.add_page_break()
heading('RESEARCH AREAS')
doc.add_paragraph('')

areas = [
    {
        'title': 'Trustworthy Generative AI',
        'icon': '🛡️',
        'description': 'Developing methods to improve the factual accuracy and reliability of large language models, with a focus on hallucination detection, grounding, and fact-checking at scale.',
        'members': 'xinyi-zhou',
    },
    {
        'title': 'Personalized & Socially Intelligent LLMs',
        'icon': '🤝',
        'description': 'Building language models that adapt to individual users and social contexts — personalizing responses, modeling user intent, and navigating complex interpersonal dynamics.',
        'members': 'xinyi-zhou',
    },
    {
        'title': 'Multimodal Machine Learning',
        'icon': '🌐',
        'description': 'Learning rich representations across text, images, and structured data to support reasoning tasks that span multiple modalities — from medical imaging to social media.',
        'members': 'xinyi-zhou, prabal-shrestha',
    },
    {
        'title': 'LLM Agents & Tool-Use',
        'icon': '🤖',
        'description': 'Studying how large language models can act as autonomous agents that plan, use external tools, and complete multi-step tasks reliably and safely.',
        'members': 'xinyi-zhou',
    },
    {
        'title': 'Applied AI in High-Stakes Domains',
        'icon': '🏥',
        'description': 'Translating research into practical systems for healthcare, education, and fact-checking — domains where accuracy, fairness, and interpretability are non-negotiable.',
        'members': 'xinyi-zhou, benjamin-lee-peterson',
    },
]

for i, a in enumerate(areas):
    item_title(a['title'])
    field('Icon', a['icon'])
    field('Description', a['description'])
    field('Members', a['members'])
    if i < len(areas) - 1:
        spacer()

# ── save ──────────────────────────────────────────────────

doc.save('CONTENT_UPDATES.docx')
print('Saved CONTENT_UPDATES.docx')
